import os
from docx import Document
from docx.oxml.ns import qn


def convert_docx_to_markdown(docx_path, output_md_path, images_dir):
    document = Document(docx_path)
    os.makedirs(images_dir, exist_ok=True)
    md_content = []

    def get_list_paragraph_level(paragraph):
        """
        Get the list level of the paragraph based on its style.
        """
        num_id = paragraph._element.xpath(".//w:numId")
        if not num_id:
            return 0
        num_id = num_id[0].get(qn("w:val"))
        level_id = paragraph._element.xpath(".//w:ilvl")
        if not level_id:
            return 0
        level_id = level_id[0].get(qn("w:val"))
        return int(level_id) + 1

    def add_paragraph_as_md(paragraph):
        """
        Convert a paragraph to its Markdown representation.
        """
        level = get_list_paragraph_level(paragraph)
        if level > 0:
            prefix = "  " * (level - 1) + "* "
        else:
            prefix = ""
        md_content.append(f"{prefix}{paragraph.text}")

    for paragraph in document.paragraphs:
        add_paragraph_as_md(paragraph)

    for rel in document.part.rels.values():
        if "image" in rel.target_ref:
            image = rel.target_part.blob
            image_name = os.path.basename(rel.target_ref)
            image_path = os.path.join(images_dir, image_name)
            with open(image_path, "wb") as img_file:
                img_file.write(image)
            md_content.append(f"![{image_name}](./images/{image_name})")
            # Insert image at the correct position in the markdown content
            for i, paragraph in enumerate(document.paragraphs):

                if image_name in paragraph.text:
                    md_content.insert(i, f"![{image_name}](./images/{image_name})")
                    break

    with open(output_md_path, "w") as md_file:
        md_file.write("\n\n".join(md_content))

    return output_md_path


def read_markdown_file(md_file_path):
    with open(md_file_path, "r") as md_file:
        return md_file.read()
